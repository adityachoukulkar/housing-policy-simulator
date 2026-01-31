#!/usr/bin/env python3
"""Create a polished DOCX for the pass-through model with clear equations."""

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import yaml


def set_base_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)


def add_equation(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Cambria Math"
    run.font.size = Pt(12)


def add_table(doc: Document, title: str, rows: list[tuple[str, str]]) -> None:
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Parameter"
    hdr_cells[1].text = "Default / Notes"
    for name, value in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = value


def load_config(path: Path) -> Dict[str, Any]:
    with path.open("r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def main() -> int:
    doc = Document()
    set_base_style(doc)

    cfg = load_config(Path("config/sim_params.yaml"))

    doc.add_heading("Structural Model Summary", level=1)

    doc.add_heading("Goal", level=2)
    doc.add_paragraph(
        "Make rental tax pass-through respond to market tightness (vacancy) and demand elasticity, "
        "rather than treating it as a fixed coefficient."
    )

    doc.add_heading("Definitions", level=2)
    items = [
        "dTax_rent: incremental effective tax rate on rental properties (e.g., 0.002 = +0.20%)",
        "Vac_t: vacancy rate at time t (proxy for market tightness)",
        "ε_d: demand elasticity (higher means renters are more price-sensitive)",
        "φ_t: pass-through rate (bounded 0 to 1)",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Pass-Through Function", level=2)
    doc.add_paragraph("We model pass-through as a bounded function of vacancy and demand elasticity:")
    add_equation(doc, "φ_t = clamp( φ0 + φ1 · (Vac_target − Vac_t) − φ2 · ε_d , 0, 1 )")

    doc.add_paragraph("Where:")
    items = [
        "φ0: base pass-through rate",
        "φ1 > 0 increases pass-through when vacancy is below target (tight market)",
        "φ2 > 0 reduces pass-through when demand is more elastic",
        "Vac_target: reference vacancy rate (e.g., 5%)",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Updated Rent Growth Equation", level=2)
    doc.add_paragraph("Replace the fixed tax term with an endogenous pass-through term:")
    add_equation(
        doc,
        "log(R_t / R_{t−1}) = a0 + a1·log(H_t / H_{t−1}) + a2·log(Pop_t / Pop_{t−1}) "
        "+ a3·(φ_t · dTax_rent) + a4·Vac_{t−1}",
    )

    doc.add_heading("Interpretation", level=2)
    items = [
        "If vacancy is low, φ_t rises and more of the tax increase is passed to renters.",
        "If demand is more elastic, φ_t falls and pass-through is muted.",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Assumptions and Limitations", level=2)
    items = [
        "Statewide aggregation masks local heterogeneity in vacancy, rents, and zoning.",
        "Vacancy is proxied using available data; if ACS is missing, a rental vacancy proxy is used.",
        "Demand elasticity is treated as a fixed parameter rather than estimated from microdata.",
        "Pass-through is bounded in [0, 1] and does not model landlord exit or conversion decisions.",
        "Upzoning is represented as a percentage uplift to completions rather than parcel-level capacity.",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Calibration Guidance (Small Data)", level=2)
    items = [
        "Fix ε_d from literature or choose a plausible range (e.g., 0.3–1.5).",
        "Set Vac_target to the long-run vacancy average (e.g., 5%).",
        "Choose φ0 as a baseline (0.3–0.6), then test sensitivity using φ1 and φ2.",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    pass_through = cfg.get("pass_through", {})
    add_table(
        doc,
        "Parameter Defaults: Pass-Through",
        [
            ("pass_through.base", str(pass_through.get("base", ""))),
            ("pass_through.vacancy_target", str(pass_through.get("vacancy_target", ""))),
            ("pass_through.vacancy_slope", str(pass_through.get("vacancy_slope", ""))),
            ("pass_through.elasticity_slope", str(pass_through.get("elasticity_slope", ""))),
            ("pass_through.demand_elasticity", str(pass_through.get("demand_elasticity", ""))),
        ],
    )

    doc.add_heading("Price Model: User-Cost with Momentum", level=2)
    doc.add_paragraph(
        "Prices are anchored to user cost and can include rent-driven momentum and drift."
    )
    add_equation(
        doc,
        "P_t = (R_t / UC_t) + λ · (ΔR_t / UC_t) ;  UC_t = r + τ + m + δ − g_R",
    )
    add_equation(
        doc,
        "P_t = P_t · exp(κ_t · g_R) · exp(drift)",
    )
    doc.add_paragraph("Where momentum decays over time:")
    add_equation(doc, "κ_t = κ_0 · exp(−decay · t)")

    user_cost = cfg.get("user_cost", {})
    add_table(
        doc,
        "Parameter Defaults: User-Cost",
        [
            ("user_cost.real_rate", str(user_cost.get("real_rate", ""))),
            ("user_cost.property_tax_base", str(user_cost.get("property_tax_base", ""))),
            ("user_cost.maintenance", str(user_cost.get("maintenance", ""))),
            ("user_cost.depreciation", str(user_cost.get("depreciation", ""))),
            ("user_cost.expected_rent_growth", str(user_cost.get("expected_rent_growth", ""))),
            ("user_cost.momentum_kappa", str(user_cost.get("momentum_kappa", ""))),
            ("user_cost.momentum_decay", str(user_cost.get("momentum_decay", ""))),
            ("user_cost.price_drift", str(user_cost.get("price_drift", ""))),
            ("user_cost.rent_capitalization_lambda", str(user_cost.get("rent_capitalization_lambda", ""))),
        ],
    )

    supply_response = cfg.get("supply_response", {})
    add_table(
        doc,
        "Parameter Defaults: Supply Response",
        [
            ("supply_response.price_elasticity", str(supply_response.get("price_elasticity", ""))),
            ("supply_response.min_multiplier", str(supply_response.get("min_multiplier", ""))),
            ("supply_response.max_multiplier", str(supply_response.get("max_multiplier", ""))),
        ],
    )

    doc.add_heading("Where This Lives in Code", level=2)
    doc.add_paragraph("Configuration parameters:")
    items = [
        "config/sim_params.yaml: pass_through.base, vacancy_target, vacancy_slope, elasticity_slope, demand_elasticity",
        "scripts/simulate.py: computes φ_t each year and applies it to the tax term",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("References / Theory Notes", level=2)
    items = [
        "Incidence intuition: pass-through rises in tighter markets and falls with more elastic demand.",
        "User-cost framing: property taxes are a cost component that can be partially passed to rents.",
        "This document provides a transparent baseline; parameters should be calibrated and sensitivity-tested.",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    out_path = Path("docs/structural_model_pass_through.docx")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"Wrote {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
