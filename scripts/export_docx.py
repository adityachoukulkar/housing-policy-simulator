#!/usr/bin/env python3
"""Export a markdown doc to .docx (minimal formatting)."""

from __future__ import annotations

from pathlib import Path

from docx import Document


def add_paragraph(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def export_md_to_docx(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    lines = md_path.read_text(encoding="utf-8").splitlines()

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif stripped.startswith("```"):
            # Skip code fence markers; render code lines plainly until next fence
            continue
        elif stripped == "":
            doc.add_paragraph("")
        else:
            add_paragraph(doc, stripped)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def main() -> int:
    md_path = Path("docs/structural_model_pass_through.md")
    docx_path = Path("docs/structural_model_pass_through.docx")
    if not md_path.exists():
        raise SystemExit(f"Missing markdown file: {md_path}")
    export_md_to_docx(md_path, docx_path)
    print(f"Wrote {docx_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
